"""
parser.py - Resume document parsing module.

Extracts raw text and structured information from PDF and DOCX files.
Uses PyMuPDF (fitz) for PDF extraction and python-docx for DOCX extraction.
"""

import re
import io
from typing import Dict, List, Optional, Any

import fitz  # PyMuPDF
from docx import Document

from utils.helpers import (
    clean_text, extract_email, extract_phone,
    extract_links, extract_name
)
from utils.constants import SECTION_PATTERNS


def parse_pdf(file_bytes: bytes) -> Dict[str, Any]:
    """
    Parse a PDF file and extract text with page-level metadata.
    
    Args:
        file_bytes: Raw bytes of the PDF file.
    
    Returns:
        Dictionary with 'text', 'pages', and 'metadata'.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        full_text = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            cleaned = clean_text(text)
            pages.append({
                "page_number": page_num + 1,
                "text": cleaned,
                "char_count": len(cleaned),
            })
            full_text.append(cleaned)
        
        combined_text = "\n\n".join(full_text)
        doc.close()
        
        return {
            "text": combined_text,
            "pages": pages,
            "page_count": len(pages),
            "format": "pdf",
            "error": None,
        }
    except Exception as e:
        return {
            "text": "",
            "pages": [],
            "page_count": 0,
            "format": "pdf",
            "error": f"Failed to parse PDF: {str(e)}",
        }


def parse_docx(file_bytes: bytes) -> Dict[str, Any]:
    """
    Parse a DOCX file and extract text.
    
    Args:
        file_bytes: Raw bytes of the DOCX file.
    
    Returns:
        Dictionary with 'text', 'pages', and 'metadata'.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        combined_text = clean_text("\n".join(paragraphs))
        
        return {
            "text": combined_text,
            "pages": [{"page_number": 1, "text": combined_text, "char_count": len(combined_text)}],
            "page_count": 1,
            "format": "docx",
            "error": None,
        }
    except Exception as e:
        return {
            "text": "",
            "pages": [],
            "page_count": 0,
            "format": "docx",
            "error": f"Failed to parse DOCX: {str(e)}",
        }


def parse_resume(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Parse a resume file (PDF or DOCX) and extract structured information.
    
    Args:
        file_bytes: Raw bytes of the file.
        filename: Original filename (used to determine format).
    
    Returns:
        Dictionary with extracted text, sections, contact info, and metadata.
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    
    if ext == 'pdf':
        parsed = parse_pdf(file_bytes)
    elif ext == 'docx':
        parsed = parse_docx(file_bytes)
    else:
        return {"error": f"Unsupported format: {ext}"}
    
    if parsed["error"]:
        return parsed
    
    text = parsed["text"]
    
    # Extract contact information
    contact_info = {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "links": extract_links(text),
    }
    
    # Detect and extract sections
    sections = detect_sections(text)
    
    # Extract skills from the text
    skills = extract_skills_from_text(text)
    
    return {
        "text": text,
        "pages": parsed["pages"],
        "page_count": parsed["page_count"],
        "format": parsed["format"],
        "contact_info": contact_info,
        "sections": sections,
        "skills": skills,
        "filename": filename,
        "error": None,
    }


def detect_sections(text: str) -> Dict[str, str]:
    """
    Detect and extract resume sections using pattern matching.
    
    Identifies section headings and extracts the content below each heading
    until the next section begins.
    
    Returns:
        Dictionary mapping section names to their content.
    """
    lines = text.split('\n')
    sections = {}
    current_section = "header"
    current_content = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_content.append("")
            continue
        
        # Check if this line is a section header
        detected_section = _identify_section_header(stripped)
        
        if detected_section:
            # Save the previous section
            if current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    sections[current_section] = content
            
            current_section = detected_section
            current_content = []
        else:
            current_content.append(stripped)
    
    # Save the last section
    if current_content:
        content = '\n'.join(current_content).strip()
        if content:
            sections[current_section] = content
    
    return sections


def _identify_section_header(line: str) -> Optional[str]:
    """
    Determine if a line is a section header.
    
    Uses pattern matching against known section heading patterns.
    A header is typically a short line (< 50 chars) that matches known patterns.
    """
    cleaned = line.lower().strip()
    
    # Remove common formatting characters
    cleaned = re.sub(r'^[#\-=*•|:]+\s*', '', cleaned)
    cleaned = re.sub(r'\s*[#\-=*•|:]+$', '', cleaned)
    cleaned = cleaned.strip()
    
    if not cleaned or len(cleaned) > 60:
        return None
    
    # Check against known patterns
    for section_name, patterns in SECTION_PATTERNS.items():
        for pattern in patterns:
            if cleaned == pattern or cleaned.startswith(pattern + ' ') or cleaned.startswith(pattern + ':'):
                return section_name
            # Also check if the cleaned line is contained in the pattern
            if pattern in cleaned and len(cleaned) < 40:
                return section_name
    
    return None


def extract_skills_from_text(text: str) -> List[str]:
    """
    Extract skills mentioned in the resume text.
    
    Uses a comprehensive skills database for matching.
    """
    from utils.constants import ALL_SKILLS
    
    text_lower = text.lower()
    found_skills = []
    
    for skill in ALL_SKILLS:
        # Use word boundary matching for short skills to avoid false positives
        if len(skill) <= 2:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill)
        else:
            if skill in text_lower:
                found_skills.append(skill)
    
    return list(set(found_skills))
